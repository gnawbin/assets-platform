"""Office 文档解析器：Word / Excel → 纯文本"""

import time
from models.parse_result import ParseResult


class DocsParser:
    """Office 文档解析器：Word（doc/docx）/ Excel（xls/xlsx）→ 纯文本"""

    async def parse(self, file_path: str, options: dict = None) -> ParseResult:
        start = time.time()
        ext = file_path.rsplit(".", 1)[-1].lower()

        try:
            if ext in ("docx", "doc"):
                raw_text = self._parse_word(file_path)
            elif ext in ("xlsx", "xls"):
                raw_text = self._parse_excel(file_path)
            else:
                raw_text = ""
        except Exception as e:
            print(f"[DocsParser] 解析失败 {ext}: {e}")
            raw_text = ""

        return ParseResult(
            file_name=file_path.rsplit("/", 1)[-1],
            file_type="document",
            raw_text=raw_text,
            images=[],  # Office 文档无需 VLM 描述
            metadata={
                "sub_type": ext,
                "parse_duration_ms": int((time.time() - start) * 1000),
            },
        )

    def _parse_word(self, file_path: str) -> str:
        """python-docx 提取段落 + 表格文字（仅支持 .docx；.doc 需先转 OOXML）"""
        from docx import Document

        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # 提取表格内容
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))

        return "\n".join(paragraphs)

    def _parse_excel(self, file_path: str) -> str:
        """openpyxl 提取所有工作表内容（仅支持 .xlsx；.xls 需先转 OOXML）"""
        from openpyxl import load_workbook

        wb = load_workbook(file_path, read_only=True, data_only=True)
        lines = []
        for sheet in wb.sheetnames:
            ws = wb[sheet]
            lines.append(f"【工作表: {sheet}】")
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) for c in row if c is not None]
                if cells:
                    lines.append(" | ".join(cells))
        wb.close()
        return "\n".join(lines)